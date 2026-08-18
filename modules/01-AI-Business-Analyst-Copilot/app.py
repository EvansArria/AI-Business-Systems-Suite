import os
import re
import csv
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime
from io import BytesIO, StringIO
from xml.sax.saxutils import escape

import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

st.set_page_config(
    page_title="AI Business Analyst Copilot",
    page_icon="🤖",
    layout="wide",
)

load_dotenv(override=True)
api_key = os.getenv("OPENAI_API_KEY")

if not api_key:
    try:
        api_key = st.secrets["OPENAI_API_KEY"]
    except Exception:
        api_key = None


def clean_inline_markdown(text):
    return text.replace("**", "").replace("__", "").replace("`", "")


def safe_filename(value, fallback="Business_Requirements"):
    cleaned = value.strip().replace(" ", "_")
    cleaned = "".join(
        c for c in cleaned if c.isalnum() or c in ("_", "-")
    )
    return cleaned or fallback


def generate_ai_text(prompt, spinner_message):
    client = OpenAI(api_key=api_key)
    with st.spinner(spinner_message):
        response = client.responses.create(
            model="gpt-4.1-mini",
            input=prompt,
        )
    return response.output_text



# =========================================================
# DASHBOARD METRIC HELPERS
# =========================================================

def count_unique_ids(text, prefix):
    if not text:
        return 0

    pattern = rf"\b{re.escape(prefix)}-\d{{3}}\b"

    return len(
        set(
            re.findall(
                pattern,
                text,
            )
        )
    )


def get_requirement_ids(text):
    if not text:
        return []

    requirement_ids = re.findall(
        r"\b(?:FR|NFR)-\d{3}\b",
        text,
    )

    return sorted(
        set(requirement_ids),
        key=lambda item: (
            item.startswith("NFR"),
            item,
        ),
    )


def parse_rtm_coverage(rtm_text):
    coverage_counts = {
        "Covered": 0,
        "Partial": 0,
        "Gap": 0,
    }

    if not rtm_text:
        return coverage_counts

    for line in rtm_text.splitlines():
        stripped = line.strip()

        if not stripped.startswith("|"):
            continue

        if "---" in stripped:
            continue

        cells = [
            cell.strip()
            for cell in stripped.strip("|").split("|")
        ]

        if len(cells) < 6:
            continue

        if cells[0] == "Requirement ID":
            continue

        status = cells[5]

        if status in coverage_counts:
            coverage_counts[status] += 1

    return coverage_counts


def calculate_traceability_score(coverage_counts):
    total = sum(
        coverage_counts.values()
    )

    if total == 0:
        return 0

    weighted_coverage = (
        coverage_counts["Covered"]
        + (
            0.5
            * coverage_counts["Partial"]
        )
    )

    return round(
        (
            weighted_coverage
            / total
        )
        * 100
    )


# =========================================================
# BUSINESS DATA ANALYSIS HELPERS
# =========================================================

def column_index_from_reference(cell_reference):
    letters = "".join(
        character
        for character in cell_reference
        if character.isalpha()
    )

    column_index = 0

    for character in letters.upper():
        column_index = (
            column_index * 26
            + ord(character)
            - ord("A")
            + 1
        )

    return max(
        column_index - 1,
        0,
    )


def make_unique_headers(raw_headers):
    headers = []
    seen = {}

    for index, raw_header in enumerate(
        raw_headers,
        start=1,
    ):
        header = str(
            raw_header
            if raw_header is not None
            else ""
        ).strip()

        if not header:
            header = f"Column_{index}"

        if header in seen:
            seen[header] += 1
            header = f"{header}_{seen[header]}"
        else:
            seen[header] = 1

        headers.append(header)

    return headers


def rows_to_records(rows):
    if not rows:
        return [], []

    headers = make_unique_headers(rows[0])
    records = []

    for raw_row in rows[1:]:
        row = list(raw_row)

        if len(row) < len(headers):
            row.extend([""] * (len(headers) - len(row)))

        row = row[:len(headers)]

        if not any(
            str(value).strip()
            for value in row
            if value is not None
        ):
            continue

        record = {
            headers[index]: "" if value is None else value
            for index, value in enumerate(row)
        }

        records.append(record)

    return headers, records


def read_csv_rows(file_bytes):
    decoded_text = None

    for encoding in (
        "utf-8-sig",
        "utf-8",
        "cp1252",
    ):
        try:
            decoded_text = file_bytes.decode(encoding)
            break
        except UnicodeDecodeError:
            continue

    if decoded_text is None:
        raise ValueError(
            "The CSV file encoding could not be read. "
            "Save it as UTF-8 CSV and try again."
        )

    reader = csv.reader(StringIO(decoded_text))
    return [row for row in reader]


def read_xlsx_rows(file_bytes):
    workbook_namespace = (
        "http://schemas.openxmlformats.org/"
        "spreadsheetml/2006/main"
    )

    relationships_namespace = (
        "http://schemas.openxmlformats.org/"
        "package/2006/relationships"
    )

    office_relationship_namespace = (
        "http://schemas.openxmlformats.org/"
        "officeDocument/2006/relationships"
    )

    with zipfile.ZipFile(BytesIO(file_bytes)) as workbook_zip:
        file_names = set(workbook_zip.namelist())

        required_files = {
            "xl/workbook.xml",
            "xl/_rels/workbook.xml.rels",
        }

        if not required_files.issubset(file_names):
            raise ValueError(
                "This does not appear to be a valid .xlsx workbook."
            )

        shared_strings = []

        if "xl/sharedStrings.xml" in file_names:
            shared_root = ET.fromstring(
                workbook_zip.read("xl/sharedStrings.xml")
            )

            for shared_item in shared_root.findall(
                f"{{{workbook_namespace}}}si"
            ):
                text_parts = [
                    text_node.text or ""
                    for text_node in shared_item.iter(
                        f"{{{workbook_namespace}}}t"
                    )
                ]
                shared_strings.append("".join(text_parts))

        workbook_root = ET.fromstring(
            workbook_zip.read("xl/workbook.xml")
        )

        first_sheet = workbook_root.find(
            f".//{{{workbook_namespace}}}sheet"
        )

        if first_sheet is None:
            raise ValueError(
                "The workbook does not contain a worksheet."
            )

        relationship_id = first_sheet.get(
            f"{{{office_relationship_namespace}}}id"
        )

        relationships_root = ET.fromstring(
            workbook_zip.read("xl/_rels/workbook.xml.rels")
        )

        worksheet_target = None

        for relationship in relationships_root.findall(
            f"{{{relationships_namespace}}}Relationship"
        ):
            if relationship.get("Id") == relationship_id:
                worksheet_target = relationship.get("Target")
                break

        if not worksheet_target:
            raise ValueError(
                "The first worksheet could not be located."
            )

        if worksheet_target.startswith("/"):
            worksheet_path = worksheet_target.lstrip("/")
        elif worksheet_target.startswith("xl/"):
            worksheet_path = worksheet_target
        else:
            worksheet_path = "xl/" + worksheet_target.lstrip("./")

        if worksheet_path not in file_names:
            raise ValueError(
                "The first worksheet data could not be read."
            )

        sheet_root = ET.fromstring(
            workbook_zip.read(worksheet_path)
        )

        parsed_rows = []

        for row_node in sheet_root.findall(
            f".//{{{workbook_namespace}}}row"
        ):
            cell_values = {}
            max_column_index = -1

            for cell_node in row_node.findall(
                f"{{{workbook_namespace}}}c"
            ):
                cell_reference = cell_node.get("r", "A1")
                column_index = column_index_from_reference(
                    cell_reference
                )
                max_column_index = max(
                    max_column_index,
                    column_index,
                )

                cell_type = cell_node.get("t")
                value = ""

                if cell_type == "inlineStr":
                    text_parts = [
                        text_node.text or ""
                        for text_node in cell_node.iter(
                            f"{{{workbook_namespace}}}t"
                        )
                    ]
                    value = "".join(text_parts)
                else:
                    value_node = cell_node.find(
                        f"{{{workbook_namespace}}}v"
                    )
                    raw_value = (
                        value_node.text
                        if value_node is not None
                        and value_node.text is not None
                        else ""
                    )

                    if cell_type == "s" and raw_value != "":
                        try:
                            value = shared_strings[int(raw_value)]
                        except (ValueError, IndexError):
                            value = raw_value
                    elif cell_type == "b":
                        value = "TRUE" if raw_value == "1" else "FALSE"
                    else:
                        value = raw_value

                cell_values[column_index] = value

            if max_column_index < 0:
                parsed_rows.append([])
                continue

            parsed_rows.append(
                [
                    cell_values.get(index, "")
                    for index in range(max_column_index + 1)
                ]
            )

        return parsed_rows


def read_uploaded_table(
    uploaded_file,
    max_rows=5000,
    max_columns=50,
):
    file_name = uploaded_file.name.lower()
    file_bytes = uploaded_file.getvalue()

    if file_name.endswith(".csv"):
        raw_rows = read_csv_rows(file_bytes)
    elif file_name.endswith(".xlsx"):
        raw_rows = read_xlsx_rows(file_bytes)
    else:
        raise ValueError(
            "Please upload a .csv or .xlsx file."
        )

    raw_rows = raw_rows[:max_rows + 1]
    raw_rows = [row[:max_columns] for row in raw_rows]

    headers, records = rows_to_records(raw_rows)

    if not headers:
        raise ValueError(
            "The uploaded file does not contain readable tabular data."
        )

    return headers, records


def parse_number(value):
    if value is None:
        return None

    if isinstance(value, (int, float)):
        return float(value)

    text = str(value).strip()

    if not text:
        return None

    is_percent = text.endswith("%")

    cleaned = (
        text
        .replace(",", "")
        .replace("$", "")
        .replace("£", "")
        .replace("€", "")
        .replace("%", "")
        .strip()
    )

    if cleaned.startswith("(") and cleaned.endswith(")"):
        cleaned = "-" + cleaned[1:-1]

    try:
        number = float(cleaned)

        if is_percent:
            number = number / 100

        return number
    except ValueError:
        return None


def create_data_profile(headers, records):
    profile_rows = []
    total_cells = len(records) * len(headers)
    missing_cells = 0
    numeric_columns = 0

    for header in headers:
        raw_values = [record.get(header, "") for record in records]

        nonblank_values = [
            value
            for value in raw_values
            if str(value).strip()
        ]

        missing_count = len(raw_values) - len(nonblank_values)
        missing_cells += missing_count

        numeric_values = [
            parsed
            for value in nonblank_values
            if (parsed := parse_number(value)) is not None
        ]

        numeric_ratio = (
            len(numeric_values) / len(nonblank_values)
            if nonblank_values
            else 0
        )

        is_numeric = bool(nonblank_values) and numeric_ratio >= 0.8
        column_type = "Numeric" if is_numeric else "Text / Categorical"

        if is_numeric:
            numeric_columns += 1

        unique_count = len(
            set(str(value) for value in nonblank_values)
        )

        profile_record = {
            "Column": header,
            "Type": column_type,
            "Non-Blank": len(nonblank_values),
            "Missing": missing_count,
            "Unique": unique_count,
        }

        if is_numeric and numeric_values:
            profile_record["Min"] = round(min(numeric_values), 2)
            profile_record["Average"] = round(
                sum(numeric_values) / len(numeric_values),
                2,
            )
            profile_record["Max"] = round(max(numeric_values), 2)
        else:
            profile_record["Min"] = ""
            profile_record["Average"] = ""
            profile_record["Max"] = ""

        profile_rows.append(profile_record)

    completeness = (
        (1 - (missing_cells / total_cells)) * 100
        if total_cells
        else 0
    )

    summary = {
        "rows": len(records),
        "columns": len(headers),
        "numeric_columns": numeric_columns,
        "missing_cells": missing_cells,
        "completeness": round(completeness, 1),
    }

    return summary, profile_rows


def compact_records_for_ai(
    headers,
    records,
    row_limit=40,
    column_limit=20,
):
    selected_headers = headers[:column_limit]
    lines = [
        " | ".join(selected_headers),
        " | ".join(["---" for _ in selected_headers]),
    ]

    for record in records[:row_limit]:
        values = []

        for header in selected_headers:
            value = str(record.get(header, ""))
            value = (
                value
                .replace("\n", " ")
                .replace("\r", " ")
                .replace("|", "/")
            )

            if len(value) > 120:
                value = value[:117] + "..."

            values.append(value)

        lines.append(" | ".join(values))

    return "\n".join(lines)


def compact_profile_for_ai(profile_rows, row_limit=30):
    lines = []

    for profile in profile_rows[:row_limit]:
        lines.append(
            f"- {profile['Column']}: "
            f"type={profile['Type']}; "
            f"nonblank={profile['Non-Blank']}; "
            f"missing={profile['Missing']}; "
            f"unique={profile['Unique']}; "
            f"min={profile['Min']}; "
            f"average={profile['Average']}; "
            f"max={profile['Max']}"
        )

    return "\n".join(lines)



# =========================================================
# KPI & VISUALIZATION HELPERS
# =========================================================

def get_numeric_columns(profile_rows):
    return [
        row["Column"]
        for row in profile_rows
        if row.get("Type") == "Numeric"
    ]


def get_dimension_columns(headers, numeric_columns):
    numeric_set = set(numeric_columns)

    return [
        header
        for header in headers
        if header not in numeric_set
    ]


def numeric_values_for_column(records, column):
    values = []

    for record in records:
        parsed = parse_number(
            record.get(column, "")
        )

        if parsed is not None:
            values.append(parsed)

    return values


def format_metric_value(column_name, value):
    if value is None:
        return "N/A"

    lowered = column_name.lower()

    if (
        "percent" in lowered
        or "percentage" in lowered
        or lowered.endswith("_pct")
    ):
        return f"{value:.1f}%"

    if (
        "csat" in lowered
        or "score" in lowered
    ):
        return f"{value:.2f}"

    if abs(value - round(value)) < 0.000001:
        return f"{int(round(value)):,}"

    return f"{value:,.2f}"


def find_column_by_keywords(headers, keyword_groups):
    normalized = {
        header: header.lower().replace(" ", "_")
        for header in headers
    }

    for keywords in keyword_groups:
        for header, normalized_header in normalized.items():
            if all(
                keyword in normalized_header
                for keyword in keywords
            ):
                return header

    return None


def build_kpi_cards(
    headers,
    records,
    numeric_columns,
):
    specifications = [
        (
            "Average SLA",
            [
                ("sla", "percent"),
                ("sla",),
            ],
            "average",
        ),
        (
            "Total Escalations",
            [
                ("escalation",),
            ],
            "sum",
        ),
        (
            "Average CSAT",
            [
                ("csat",),
                ("customer", "satisfaction"),
            ],
            "average",
        ),
        (
            "Average Backlog",
            [
                ("backlog",),
            ],
            "average",
        ),
        (
            "Total Tickets Opened",
            [
                ("ticket", "opened"),
                ("tickets", "opened"),
            ],
            "sum",
        ),
        (
            "Avg Resolution Hours",
            [
                ("resolution", "hour"),
                ("resolution", "time"),
            ],
            "average",
        ),
    ]

    cards = []
    used_columns = set()

    for label, keyword_groups, aggregation in specifications:
        column = find_column_by_keywords(
            numeric_columns,
            keyword_groups,
        )

        if (
            not column
            or column in used_columns
        ):
            continue

        values = numeric_values_for_column(
            records,
            column,
        )

        if not values:
            continue

        if aggregation == "sum":
            value = sum(values)
        else:
            value = sum(values) / len(values)

        cards.append(
            {
                "label": label,
                "column": column,
                "value": value,
            }
        )

        used_columns.add(column)

        if len(cards) == 4:
            break

    if len(cards) < 4:
        for column in numeric_columns:
            if column in used_columns:
                continue

            values = numeric_values_for_column(
                records,
                column,
            )

            if not values:
                continue

            cards.append(
                {
                    "label": f"Average {column}",
                    "column": column,
                    "value": sum(values) / len(values),
                }
            )

            used_columns.add(column)

            if len(cards) == 4:
                break

    return cards


def aggregate_records(
    records,
    dimension,
    metric=None,
    aggregation="Average",
):
    groups = {}

    for record in records:
        raw_dimension = record.get(
            dimension,
            "",
        )

        dimension_value = (
            str(raw_dimension).strip()
            or "(Blank)"
        )

        if dimension_value not in groups:
            groups[dimension_value] = {
                "count": 0,
                "values": [],
            }

        groups[dimension_value]["count"] += 1

        if metric:
            parsed = parse_number(
                record.get(
                    metric,
                    "",
                )
            )

            if parsed is not None:
                groups[dimension_value]["values"].append(
                    parsed
                )

    output = []

    for dimension_value, group in groups.items():

        if aggregation == "Count":
            result_value = group["count"]

        else:
            values = group["values"]

            if not values:
                continue

            if aggregation == "Sum":
                result_value = sum(values)
            elif aggregation == "Minimum":
                result_value = min(values)
            elif aggregation == "Maximum":
                result_value = max(values)
            else:
                result_value = sum(values) / len(values)

        output.append(
            {
                dimension: dimension_value,
                "Value": round(
                    result_value,
                    3,
                ),
            }
        )

    output.sort(
        key=lambda row: row["Value"],
        reverse=True,
    )

    return output


def parse_date_value(value):
    text = str(value).strip()

    if not text:
        return None

    formats = (
        "%Y-%m-%d",
        "%Y/%m/%d",
        "%m/%d/%Y",
        "%m-%d-%Y",
        "%Y-%m-%d %H:%M:%S",
        "%Y-%m-%dT%H:%M:%S",
    )

    try:
        return datetime.fromisoformat(
            text.replace(
                "Z",
                "+00:00",
            )
        )
    except ValueError:
        pass

    for format_string in formats:
        try:
            return datetime.strptime(
                text,
                format_string,
            )
        except ValueError:
            continue

    return None


def detect_date_columns(headers, records):
    detected = []

    for header in headers:
        lowered = header.lower()

        if "date" in lowered:
            detected.append(header)
            continue

        sample_values = [
            record.get(header, "")
            for record in records[:50]
            if str(
                record.get(
                    header,
                    "",
                )
            ).strip()
        ]

        if not sample_values:
            continue

        parsed_count = sum(
            1
            for value in sample_values
            if parse_date_value(value) is not None
        )

        if (
            parsed_count
            / len(sample_values)
            >= 0.8
        ):
            detected.append(header)

    return detected


def aggregate_time_series(
    records,
    date_column,
    metric,
    aggregation="Average",
):
    grouped = {}

    for record in records:
        parsed_date = parse_date_value(
            record.get(
                date_column,
                "",
            )
        )

        parsed_metric = parse_number(
            record.get(
                metric,
                "",
            )
        )

        if (
            parsed_date is None
            or parsed_metric is None
        ):
            continue

        date_key = parsed_date.date().isoformat()

        grouped.setdefault(
            date_key,
            [],
        ).append(
            parsed_metric
        )

    output = []

    for date_key in sorted(grouped):
        values = grouped[date_key]

        if aggregation == "Sum":
            result_value = sum(values)
        elif aggregation == "Minimum":
            result_value = min(values)
        elif aggregation == "Maximum":
            result_value = max(values)
        else:
            result_value = sum(values) / len(values)

        output.append(
            {
                date_column: date_key,
                "Value": round(
                    result_value,
                    3,
                ),
            }
        )

    return output



# =========================================================
# COMPLETE PROJECT PACKAGE HELPERS
# =========================================================

def build_project_package_text(
    company,
    project_name,
    requirements,
    user_stories,
    test_cases,
    rtm,
    executive_analysis,
    data_analysis,
    data_filename,
):
    coverage_counts = parse_rtm_coverage(
        rtm or ""
    )

    traceability_score = (
        calculate_traceability_score(
            coverage_counts
        )
        if rtm
        else 0
    )

    sections = [
        "# AI Business Analyst Copilot — Complete Project Package",
        "",
        f"**Company:** {company or 'Not provided'}",
        f"**Project:** {project_name or 'Business Analysis Project'}",
        "",
        "## Package Overview",
        "",
        (
            "This package consolidates the business-analysis artifacts "
            "generated during the project workflow."
        ),
        "",
        "### Artifact Status",
        "",
        (
            f"- Business Requirements Document: "
            f"{'Included' if requirements else 'Not generated'}"
        ),
        (
            f"- Jira User Stories: "
            f"{'Included' if user_stories else 'Not generated'}"
        ),
        (
            f"- QA Test Cases: "
            f"{'Included' if test_cases else 'Not generated'}"
        ),
        (
            f"- Requirements Traceability Matrix: "
            f"{'Included' if rtm else 'Not generated'}"
        ),
        (
            f"- Executive Analysis: "
            f"{'Included' if executive_analysis else 'Not generated'}"
        ),
        (
            f"- Business Data Analysis: "
            f"{'Included' if data_analysis else 'Not generated'}"
        ),
        "",
    ]

    if rtm:
        sections.extend(
            [
                "### Traceability Snapshot",
                "",
                (
                    f"- Covered Requirements: "
                    f"{coverage_counts['Covered']}"
                ),
                (
                    f"- Partial Coverage: "
                    f"{coverage_counts['Partial']}"
                ),
                (
                    f"- Coverage Gaps: "
                    f"{coverage_counts['Gap']}"
                ),
                (
                    f"- Traceability Score: "
                    f"{traceability_score}%"
                ),
                "",
            ]
        )

    sections.extend(
        [
            "## Table of Contents",
            "",
            "1. Business Requirements Document",
            "2. Jira-Ready User Stories",
            "3. QA Test Cases",
            "4. Requirements Traceability Matrix",
            "5. Executive Project Analysis",
            "6. Business Data Analysis",
            "",
        ]
    )

    artifact_sections = [
        (
            "## 1. Business Requirements Document",
            requirements,
            "The BRD has not been generated.",
        ),
        (
            "## 2. Jira-Ready User Stories",
            user_stories,
            "Jira user stories have not been generated.",
        ),
        (
            "## 3. QA Test Cases",
            test_cases,
            "QA test cases have not been generated.",
        ),
        (
            "## 4. Requirements Traceability Matrix",
            rtm,
            "The traceability matrix has not been generated.",
        ),
        (
            "## 5. Executive Project Analysis",
            executive_analysis,
            "The executive analysis has not been generated.",
        ),
    ]

    for heading, content, missing_message in artifact_sections:
        sections.extend(
            [
                heading,
                "",
                content or missing_message,
                "",
            ]
        )

    sections.extend(
        [
            "## 6. Business Data Analysis",
            "",
        ]
    )

    if data_analysis:
        if data_filename:
            sections.extend(
                [
                    f"**Source File:** {data_filename}",
                    "",
                ]
            )

        sections.extend(
            [
                data_analysis,
                "",
            ]
        )
    else:
        sections.extend(
            [
                (
                    "Business data analysis was not included in this "
                    "package."
                ),
                "",
            ]
        )

    sections.extend(
        [
            "## Package Notes",
            "",
            (
                "- AI-generated artifacts should be reviewed and approved "
                "by appropriate business, product, technical, and QA "
                "stakeholders before implementation."
            ),
            (
                "- Traceability, priorities, risks, and recommendations "
                "reflect the information provided to the application."
            ),
            (
                "- This package does not constitute formal project approval "
                "or production authorization."
            ),
        ]
    )

    return "\n".join(
        sections
    )


def create_project_package_zip(
    package_text,
    company,
    project_name,
    artifacts,
):
    buffer = BytesIO()

    base_name = safe_filename(
        project_name,
        fallback="Business_Analysis_Project",
    )

    package_word = create_word_document(
        "Complete Business Analysis Project Package",
        package_text,
        company,
        project_name,
    )

    package_pdf = create_pdf_document(
        "Complete Business Analysis Project Package",
        package_text,
        company,
        project_name,
    )

    with zipfile.ZipFile(
        buffer,
        mode="w",
        compression=zipfile.ZIP_DEFLATED,
    ) as package_zip:

        package_zip.writestr(
            f"{base_name}_Complete_Project_Package.md",
            package_text,
        )

        package_zip.writestr(
            f"{base_name}_Complete_Project_Package.docx",
            package_word.getvalue(),
        )

        package_zip.writestr(
            f"{base_name}_Complete_Project_Package.pdf",
            package_pdf.getvalue(),
        )

        for artifact_name, artifact_content in artifacts:
            if not artifact_content:
                continue

            package_zip.writestr(
                f"{base_name}_{artifact_name}.md",
                artifact_content,
            )

    buffer.seek(0)

    return buffer


def create_word_document(document_title, content_text, company, project_name):
    buffer = BytesIO()
    document = Document()
    document.add_heading(document_title, level=0)
    document.add_paragraph(f"Company: {company}")
    document.add_paragraph(f"Project: {project_name}")
    document.add_paragraph("")

    for line in content_text.splitlines():
        cleaned_line = line.strip()

        if not cleaned_line:
            document.add_paragraph("")
            continue

        if cleaned_line.startswith("### "):
            document.add_heading(
                clean_inline_markdown(cleaned_line[4:]), level=3
            )
        elif cleaned_line.startswith("## "):
            document.add_heading(
                clean_inline_markdown(cleaned_line[3:]), level=2
            )
        elif cleaned_line.startswith("# "):
            document.add_heading(
                clean_inline_markdown(cleaned_line[2:]), level=1
            )
        elif cleaned_line.startswith("- "):
            document.add_paragraph(
                clean_inline_markdown(cleaned_line[2:]),
                style="List Bullet",
            )
        else:
            document.add_paragraph(clean_inline_markdown(cleaned_line))

    document.save(buffer)
    buffer.seek(0)
    return buffer


def create_pdf_document(document_title, content_text, company, project_name):
    buffer = BytesIO()
    pdf = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()
    story = [
        Paragraph(escape(document_title), styles["Title"]),
        Spacer(1, 12),
        Paragraph(f"<b>Company:</b> {escape(company)}", styles["Normal"]),
        Paragraph(f"<b>Project:</b> {escape(project_name)}", styles["Normal"]),
        Spacer(1, 18),
    ]

    for line in content_text.splitlines():
        cleaned_line = line.strip()

        if not cleaned_line:
            story.append(Spacer(1, 8))
            continue

        if cleaned_line.startswith("### "):
            story.append(
                Paragraph(
                    escape(clean_inline_markdown(cleaned_line[4:])),
                    styles["Heading3"],
                )
            )
        elif cleaned_line.startswith("## "):
            story.append(
                Paragraph(
                    escape(clean_inline_markdown(cleaned_line[3:])),
                    styles["Heading2"],
                )
            )
        elif cleaned_line.startswith("# "):
            story.append(
                Paragraph(
                    escape(clean_inline_markdown(cleaned_line[2:])),
                    styles["Heading1"],
                )
            )
        elif cleaned_line.startswith("- "):
            story.append(
                Paragraph(
                    "&bull; " + escape(clean_inline_markdown(cleaned_line[2:])),
                    styles["Normal"],
                )
            )
        else:
            story.append(
                Paragraph(
                    escape(clean_inline_markdown(cleaned_line)),
                    styles["Normal"],
                )
            )

        story.append(Spacer(1, 6))

    pdf.build(story)
    buffer.seek(0)
    return buffer


def show_download_buttons(
    document_title,
    content_text,
    company,
    project_name,
    filename_suffix,
):
    base_name = safe_filename(
        project_name,
        fallback="Business_Analysis_Project",
    )

    word_file = create_word_document(
        document_title,
        content_text,
        company,
        project_name,
    )

    pdf_file = create_pdf_document(
        document_title,
        content_text,
        company,
        project_name,
    )

    col1, col2, col3 = st.columns(3)

    with col1:
        st.download_button(
            label="📘 Download Word",
            data=word_file,
            file_name=f"{base_name}_{filename_suffix}.docx",
            mime=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            use_container_width=True,
        )

    with col2:
        st.download_button(
            label="📕 Download PDF",
            data=pdf_file,
            file_name=f"{base_name}_{filename_suffix}.pdf",
            mime="application/pdf",
            use_container_width=True,
        )

    with col3:
        st.download_button(
            label="📝 Download Markdown",
            data=content_text,
            file_name=f"{base_name}_{filename_suffix}.md",
            mime="text/markdown",
            use_container_width=True,
        )


default_state = {
    "generated_requirements": None,
    "generated_company": "",
    "generated_project_name": "",
    "generated_user_stories": None,
    "generated_test_cases": None,
    "generated_rtm": None,
    "generated_executive_analysis": None,
    "generated_data_analysis": None,
    "generated_data_filename": "",
    "generated_project_package": None,
}

for key, value in default_state.items():
    if key not in st.session_state:
        st.session_state[key] = value


st.title("🤖 AI Business Analyst Copilot")
st.caption(
    "Generate professional Business Analysis documentation using AI."
)
st.divider()

st.subheader("📋 Business Requirements Generator")

company = st.text_input("Company")
project_name = st.text_input("Project Name")
problem = st.text_area("Business Problem", height=120)
goal = st.text_area("Business Goal", height=120)
stakeholders = st.text_area("Stakeholders (one per line)", height=100)
constraints = st.text_area("Business Constraints", height=100)


if st.button(
    "✨ Generate Business Requirements",
    use_container_width=True,
):
    required_fields_missing = (
        not company.strip()
        or not project_name.strip()
        or not problem.strip()
        or not goal.strip()
    )

    if required_fields_missing:
        st.warning(
            "Please complete Company, Project Name, "
            "Business Problem, and Business Goal."
        )

    elif not api_key:
        st.error(
            "The OpenAI API key was not found. "
            "Check your .env file locally or Streamlit Secrets in the live app."
        )

    else:
        brd_prompt = f'''
You are a senior Business Systems Analyst.

Create a professional Business Requirements Document using the information below.

Company:
{company}

Project Name:
{project_name}

Business Problem:
{problem}

Business Goal:
{goal}

Stakeholders:
{stakeholders or "Not provided"}

Business Constraints:
{constraints or "Not provided"}

Create the following sections:

1. Executive Summary
2. Business Problem Statement
3. Business Objectives
4. Project Scope
5. Stakeholder Analysis
6. Functional Requirements
7. Non-Functional Requirements
8. User Stories
9. Acceptance Criteria
10. Assumptions
11. Dependencies
12. Risks and Mitigation Strategies
13. Success Metrics

Requirements:

- Use clear, professional enterprise business-analysis language.
- Make the requirements specific to the information provided.
- Number functional requirements using FR-001, FR-002, FR-003, etc.
- Number non-functional requirements using NFR-001, NFR-002, NFR-003, etc.
- Write user stories using:
  "As a [user], I want [capability], so that [business benefit]."
- Make acceptance criteria measurable wherever possible.
- Clearly distinguish assumptions from confirmed requirements.
- Do not invent highly specific company facts that were not provided.
'''

        try:
            requirements = generate_ai_text(
                brd_prompt,
                "Generating business requirements with AI...",
            )

            if not requirements:
                st.error(
                    "OpenAI returned an empty response. Please try again."
                )
            else:
                st.session_state["generated_requirements"] = requirements
                st.session_state["generated_company"] = company
                st.session_state["generated_project_name"] = project_name
                st.session_state["generated_user_stories"] = None
                st.session_state["generated_test_cases"] = None
                st.session_state["generated_rtm"] = None
                st.session_state["generated_executive_analysis"] = None

        except Exception as error:
            st.error(f"OpenAI connection error: {error}")


if st.session_state["generated_requirements"]:
    requirements = st.session_state["generated_requirements"]
    generated_company = st.session_state["generated_company"]
    generated_project_name = st.session_state["generated_project_name"]

    st.success("Business Requirements Generated")
    st.divider()
    st.header("📄 Generated Business Requirements Document")
    st.markdown(requirements)

    st.subheader("📥 Export Business Requirements")

    show_download_buttons(
        "Business Requirements Document",
        requirements,
        generated_company,
        generated_project_name,
        "BRD",
    )

    st.divider()
    st.header("📋 Jira-Ready User Stories")
    st.caption(
        "Generate implementation-ready user stories from the approved "
        "Business Requirements Document."
    )

    if st.button(
        "🧩 Generate Jira User Stories",
        use_container_width=True,
    ):
        if not api_key:
            st.error("The OpenAI API key was not found.")
        else:
            stories_prompt = f'''
You are a senior Business Analyst and Product Owner.

Using ONLY the Business Requirements Document below, create a set of
Jira-ready user stories.

Company:
{generated_company}

Project:
{generated_project_name}

BUSINESS REQUIREMENTS DOCUMENT:

{requirements}

Create enough user stories to cover the major functional requirements
without creating unnecessary duplicates.

Use this exact Markdown structure for every story:

## US-001 — [Concise Jira Summary]

- Epic: [Logical epic name]
- Priority: [High, Medium, or Low]
- Story Points: [1, 2, 3, 5, or 8]
- Linked Requirements: [FR-### and/or NFR-###]

### User Story

As a [specific user/persona], I want [capability], so that [business benefit].

### Acceptance Criteria

1. Given [context], when [action], then [measurable outcome].
2. Given [context], when [action], then [measurable outcome].
3. Given [context], when [action], then [measurable outcome].

### Business Value

[One concise sentence explaining why the story matters.]

Rules:

- Number stories sequentially: US-001, US-002, etc.
- Use requirements traceability wherever possible.
- Do not invent requirements that are not supported by the BRD.
- Write acceptance criteria in Given/When/Then format.
- Keep Jira summaries concise and action-oriented.
- Assign realistic story points based on relative complexity.
'''

            try:
                user_stories = generate_ai_text(
                    stories_prompt,
                    "Generating Jira-ready user stories...",
                )

                if not user_stories:
                    st.error("OpenAI returned an empty response.")
                else:
                    st.session_state["generated_user_stories"] = user_stories
                    st.session_state["generated_test_cases"] = None
                    st.session_state["generated_rtm"] = None
                    st.session_state["generated_executive_analysis"] = None

            except Exception as error:
                st.error(f"User story generation error: {error}")


    if st.session_state["generated_user_stories"]:
        user_stories = st.session_state["generated_user_stories"]

        st.success("Jira User Stories Generated")
        st.markdown(user_stories)

        st.subheader("📥 Export Jira User Stories")

        show_download_buttons(
            "Jira-Ready User Stories",
            user_stories,
            generated_company,
            generated_project_name,
            "Jira_User_Stories",
        )

        st.divider()
        st.header("🧪 AI-Generated QA Test Cases")
        st.caption(
            "Generate traceable functional and negative test cases "
            "from the BRD and Jira stories."
        )

        if st.button(
            "🔬 Generate QA Test Cases",
            use_container_width=True,
        ):
            if not api_key:
                st.error("The OpenAI API key was not found.")
            else:
                test_prompt = f'''
You are a senior QA Analyst working with a Business Analyst and Product Owner.

Create professional QA test cases using ONLY the Business Requirements
Document and Jira user stories provided below.

Company:
{generated_company}

Project:
{generated_project_name}

BUSINESS REQUIREMENTS DOCUMENT:

{requirements}

JIRA USER STORIES:

{user_stories}

Create test coverage for the important business flows, validation rules,
acceptance criteria, and major negative scenarios.

Use this exact Markdown structure for every test case:

## TC-001 — [Concise Test Case Title]

- Linked User Story: [US-###]
- Linked Requirement: [FR-### and/or NFR-###]
- Priority: [High, Medium, or Low]
- Test Type: [Functional, Negative, Integration, Security, or Performance]

### Objective

[What this test validates.]

### Preconditions

- [Required starting condition]

### Test Steps

1. [Action]
2. [Action]
3. [Action]

### Expected Result

[Specific measurable expected outcome.]

Rules:

- Number test cases sequentially: TC-001, TC-002, etc.
- Include positive and negative coverage where appropriate.
- Preserve traceability to user stories and requirements.
- Do not invent functionality not supported by the source documents.
- Make steps executable by a QA tester.
- Keep expected results objective and testable.
'''

                try:
                    test_cases = generate_ai_text(
                        test_prompt,
                        "Generating QA test cases...",
                    )

                    if not test_cases:
                        st.error("OpenAI returned an empty response.")
                    else:
                        st.session_state["generated_test_cases"] = test_cases
                        st.session_state["generated_rtm"] = None
                        st.session_state["generated_executive_analysis"] = None

                except Exception as error:
                    st.error(f"Test case generation error: {error}")


        if st.session_state["generated_test_cases"]:
            test_cases = st.session_state["generated_test_cases"]

            st.success("QA Test Cases Generated")
            st.markdown(test_cases)

            st.subheader("📥 Export QA Test Cases")

            show_download_buttons(
                "QA Test Cases",
                test_cases,
                generated_company,
                generated_project_name,
                "QA_Test_Cases",
            )

            st.divider()
            st.header("🔗 Requirements Traceability Matrix")
            st.caption(
                "Map each functional and non-functional requirement "
                "to its Jira user stories and QA test cases."
            )

            if st.button(
                "🧭 Generate Traceability Matrix",
                use_container_width=True,
            ):
                if not api_key:
                    st.error("The OpenAI API key was not found.")
                else:
                    rtm_prompt = f'''
You are a senior Business Systems Analyst responsible for requirements
governance, quality assurance, and traceability.

Using ONLY the source artifacts below, create a professional Requirements
Traceability Matrix (RTM).

Company:
{generated_company}

Project:
{generated_project_name}

BUSINESS REQUIREMENTS DOCUMENT:

{requirements}

JIRA USER STORIES:

{user_stories}

QA TEST CASES:

{test_cases}

Your job is to map every functional requirement (FR-###) and every
non-functional requirement (NFR-###) to the user stories and test cases
that validate it.

Return the RTM as a Markdown table using EXACTLY these columns:

| Requirement ID | Requirement Summary | Priority | User Story ID(s) | Test Case ID(s) | Coverage Status |
|---|---|---|---|---|---|

Rules:

- Include one row for every FR-### and NFR-### found in the BRD.
- Do not invent requirement IDs, user story IDs, or test case IDs.
- Requirement Summary must be concise.
- Priority must be High, Medium, or Low based on the business impact
  supported by the source artifacts.
- If multiple user stories or test cases apply, separate the IDs with commas.
- If no user story is mapped, write: Not mapped
- If no test case is mapped, write: Not mapped
- Coverage Status must be exactly one of: Covered, Partial, Gap
- Covered means the requirement has at least one mapped user story AND at
  least one mapped test case.
- Partial means it has either a mapped user story OR a mapped test case,
  but not both.
- Gap means neither is mapped.

After the table, add:

## Traceability Summary

- Total Requirements: [number]
- Covered: [number]
- Partial: [number]
- Gaps: [number]

## Recommended Actions

Provide a short prioritized list of actions for any Partial or Gap items.

Do not create new requirements merely to improve coverage. Report actual
traceability gaps honestly.
'''

                    try:
                        rtm = generate_ai_text(
                            rtm_prompt,
                            "Generating requirements traceability matrix...",
                        )

                        if not rtm:
                            st.error("OpenAI returned an empty response.")
                        else:
                            st.session_state["generated_rtm"] = rtm
                            st.session_state["generated_executive_analysis"] = None

                    except Exception as error:
                        st.error(
                            f"Traceability matrix generation error: {error}"
                        )

            if st.session_state["generated_rtm"]:
                rtm = st.session_state["generated_rtm"]

                st.success("Requirements Traceability Matrix Generated")
                st.markdown(rtm)

                st.subheader("📥 Export Traceability Matrix")

                show_download_buttons(
                    "Requirements Traceability Matrix",
                    rtm,
                    generated_company,
                    generated_project_name,
                    "Requirements_Traceability_Matrix",
                )

                # =============================================
                # EXECUTIVE PROJECT DASHBOARD
                # =============================================

                st.divider()

                st.header(
                    "📊 Executive Project Dashboard"
                )

                st.caption(
                    "A deterministic view of delivery artifacts, "
                    "requirements coverage, and traceability health."
                )

                requirement_ids = get_requirement_ids(
                    requirements
                )

                requirement_count = len(
                    requirement_ids
                )

                user_story_count = count_unique_ids(
                    user_stories,
                    "US",
                )

                test_case_count = count_unique_ids(
                    test_cases,
                    "TC",
                )

                coverage_counts = parse_rtm_coverage(
                    rtm
                )

                traceability_score = (
                    calculate_traceability_score(
                        coverage_counts
                    )
                )

                metric1, metric2, metric3, metric4 = (
                    st.columns(4)
                )

                with metric1:
                    st.metric(
                        "Requirements",
                        requirement_count,
                    )

                with metric2:
                    st.metric(
                        "User Stories",
                        user_story_count,
                    )

                with metric3:
                    st.metric(
                        "QA Test Cases",
                        test_case_count,
                    )

                with metric4:
                    st.metric(
                        "Traceability Score",
                        f"{traceability_score}%",
                    )

                coverage1, coverage2, coverage3 = (
                    st.columns(3)
                )

                with coverage1:
                    st.metric(
                        "Covered",
                        coverage_counts["Covered"],
                    )

                with coverage2:
                    st.metric(
                        "Partial",
                        coverage_counts["Partial"],
                    )

                with coverage3:
                    st.metric(
                        "Gaps",
                        coverage_counts["Gap"],
                    )

                st.progress(
                    traceability_score / 100
                )

                st.caption(
                    "Traceability Score formula: "
                    "Covered requirements receive full credit, "
                    "Partial requirements receive half credit, "
                    "and Gap requirements receive no credit."
                )

                if coverage_counts["Gap"] > 0:
                    st.warning(
                        "Traceability gaps remain. Review the RTM "
                        "before treating the project as implementation-ready."
                    )
                elif coverage_counts["Partial"] > 0:
                    st.info(
                        "No complete traceability gaps were detected, "
                        "but some requirements have partial coverage."
                    )
                else:
                    st.success(
                        "All requirements represented in the RTM "
                        "have full story and test-case coverage."
                    )


                # =============================================
                # MOSCOW PRIORITIZATION + RISK ANALYSIS
                # =============================================

                st.divider()

                st.header(
                    "🎯 MoSCoW Prioritization & Risk Analysis"
                )

                st.caption(
                    "Create an executive-ready prioritization, "
                    "risk register, readiness assessment, "
                    "and recommended next actions."
                )

                if st.button(
                    "📈 Generate Executive Analysis",
                    use_container_width=True,
                ):

                    if not api_key:

                        st.error(
                            "The OpenAI API key was not found."
                        )

                    else:

                        executive_prompt = f"""
You are a senior Business Systems Analyst,
Product Owner, and delivery governance advisor.

Create an executive project analysis using
ONLY the source artifacts below.

Company:
{generated_company}

Project:
{generated_project_name}

CURRENT TRACEABILITY METRICS:

- Requirements: {requirement_count}
- User Stories: {user_story_count}
- QA Test Cases: {test_case_count}
- Covered Requirements: {coverage_counts["Covered"]}
- Partial Requirements: {coverage_counts["Partial"]}
- Gap Requirements: {coverage_counts["Gap"]}
- Traceability Score: {traceability_score}%

BUSINESS REQUIREMENTS DOCUMENT:

{requirements}

JIRA USER STORIES:

{user_stories}

QA TEST CASES:

{test_cases}

REQUIREMENTS TRACEABILITY MATRIX:

{rtm}

Create the following sections.

## Executive Summary

Provide a concise executive-level summary of
the project, business objective, delivery
position, and most important concerns.

## MoSCoW Prioritization

Return a Markdown table using EXACTLY these columns:

| Requirement ID | MoSCoW Category | Business Rationale |
|---|---|---|

Rules:

- Include every FR-### and NFR-### from the BRD.
- Category must be exactly one of:
  Must Have
  Should Have
  Could Have
  Won't Have This Release
- Do not invent requirement IDs.
- Base the category on business criticality,
  dependencies, risk, and stated project goals.
- "Won't Have This Release" means deferred,
  not rejected.

## Risk Register

Return a Markdown table using EXACTLY these columns:

| Risk ID | Risk Description | Probability | Impact | Score | Rating | Mitigation | Owner Role |
|---|---|---|---|---|---|---|---|

Rules:

- Number risks RISK-001, RISK-002, etc.
- Probability must be an integer from 1 to 5.
- Impact must be an integer from 1 to 5.
- Score must equal Probability multiplied by Impact.
- Rating must use:
  Low = 1-5
  Medium = 6-10
  High = 11-15
  Critical = 16-25
- Identify only risks supported by the source
  artifacts or reasonable delivery risks directly
  implied by them.
- Owner Role must be a role, not a person's name.
- Mitigation must be actionable.

## Delivery Readiness Assessment

Use this exact structure:

- Readiness: [Green, Amber, or Red]
- Traceability Score: {traceability_score}%
- Key Strengths: [concise statement]
- Key Constraints: [concise statement]
- Decision Basis: [concise explanation]

Guidance:

- Green means the artifacts show strong coverage
  with no material unresolved delivery blockers.
- Amber means implementation may proceed only
  after identified gaps, risks, or dependencies
  are addressed.
- Red means major unresolved gaps or critical
  risks make implementation premature.

## Recommended Next Actions

Provide 3 to 7 prioritized actions.
Start each action with:
1.
2.
3.
and continue sequentially.

Important:

- Do not invent company facts.
- Do not hide traceability gaps.
- Do not claim that AI-generated analysis is
  formal business approval.
- Clearly distinguish delivery recommendations
  from confirmed project decisions.
"""

                        try:

                            executive_analysis = (
                                generate_ai_text(
                                    executive_prompt,
                                    (
                                        "Generating MoSCoW priorities, "
                                        "risk analysis, and executive "
                                        "readiness assessment..."
                                    ),
                                )
                            )

                            if not executive_analysis:

                                st.error(
                                    "OpenAI returned "
                                    "an empty response."
                                )

                            else:

                                st.session_state[
                                    "generated_executive_analysis"
                                ] = executive_analysis

                        except Exception as error:

                            st.error(
                                "Executive analysis "
                                f"generation error: {error}"
                            )


                # =============================================
                # DISPLAY EXECUTIVE ANALYSIS
                # =============================================

                if (
                    st.session_state[
                        "generated_executive_analysis"
                    ]
                ):

                    executive_analysis = (
                        st.session_state[
                            "generated_executive_analysis"
                        ]
                    )

                    st.success(
                        "Executive Analysis Generated"
                    )

                    st.markdown(
                        executive_analysis
                    )

                    st.subheader(
                        "📥 Export Executive Analysis"
                    )

                    show_download_buttons(
                        (
                            "Executive Project Analysis - "
                            "MoSCoW Prioritization and Risk Register"
                        ),
                        executive_analysis,
                        generated_company,
                        generated_project_name,
                        "Executive_Project_Analysis",
                    )


# =========================================================
# AI BUSINESS DATA ANALYSIS
# =========================================================

st.divider()

st.header(
    "📈 AI Business Data Analysis"
)

st.caption(
    "Upload a CSV or Excel .xlsx file to profile "
    "business data and generate an AI-assisted "
    "operational analysis."
)

uploaded_business_file = st.file_uploader(
    "Upload business data",
    type=[
        "csv",
        "xlsx",
    ],
    max_upload_size=10,
    help=(
        "Supported formats: CSV and Excel .xlsx. "
        "The analysis reads up to 5,000 rows and "
        "50 columns per upload."
    ),
    key="business_data_uploader",
)

if uploaded_business_file is not None:

    try:

        data_headers, data_records = read_uploaded_table(
            uploaded_business_file
        )

        data_summary, data_profile = create_data_profile(
            data_headers,
            data_records,
        )

        st.success(
            f"Loaded {uploaded_business_file.name}: "
            f"{data_summary['rows']} rows and "
            f"{data_summary['columns']} columns."
        )

        data_metric1, data_metric2, data_metric3, data_metric4 = (
            st.columns(4)
        )

        with data_metric1:
            st.metric(
                "Rows",
                data_summary["rows"],
            )

        with data_metric2:
            st.metric(
                "Columns",
                data_summary["columns"],
            )

        with data_metric3:
            st.metric(
                "Numeric Columns",
                data_summary["numeric_columns"],
            )

        with data_metric4:
            st.metric(
                "Data Completeness",
                f"{data_summary['completeness']}%",
            )

        st.subheader(
            "🔎 Data Preview"
        )

        preview_records = data_records[:25]

        if preview_records:
            st.dataframe(
                preview_records,
                use_container_width=True,
                hide_index=True,
            )
        else:
            st.info(
                "The file contains headers but no data rows."
            )

        st.subheader(
            "🧾 Data Profile"
        )

        st.dataframe(
            data_profile,
            use_container_width=True,
            hide_index=True,
        )

        # =================================================
        # INTERACTIVE KPI & VISUALIZATION DASHBOARD
        # =================================================

        st.divider()

        st.subheader(
            "📊 Interactive KPI Dashboard"
        )

        st.caption(
            "Explore operational KPIs and build "
            "interactive visual summaries from "
            "the uploaded dataset."
        )

        numeric_columns = get_numeric_columns(
            data_profile
        )

        dimension_columns = get_dimension_columns(
            data_headers,
            numeric_columns,
        )

        kpi_cards = build_kpi_cards(
            data_headers,
            data_records,
            numeric_columns,
        )

        if kpi_cards:

            kpi_columns = st.columns(
                len(kpi_cards)
            )

            for index, card in enumerate(
                kpi_cards
            ):
                with kpi_columns[index]:
                    st.metric(
                        card["label"],
                        format_metric_value(
                            card["column"],
                            card["value"],
                        ),
                    )

        else:

            st.info(
                "No numeric KPI fields were "
                "detected in this dataset."
            )

        if (
            dimension_columns
            and numeric_columns
        ):

            st.markdown(
                "#### Category Performance"
            )

            chart_col1, chart_col2, chart_col3 = (
                st.columns(3)
            )

            with chart_col1:

                selected_dimension = st.selectbox(
                    "Group by",
                    options=dimension_columns,
                    key="dashboard_dimension",
                )

            with chart_col2:

                selected_metric = st.selectbox(
                    "Metric",
                    options=numeric_columns,
                    key="dashboard_metric",
                )

            with chart_col3:

                selected_aggregation = st.selectbox(
                    "Aggregation",
                    options=[
                        "Average",
                        "Sum",
                        "Minimum",
                        "Maximum",
                        "Count",
                    ],
                    key="dashboard_aggregation",
                )

            category_chart_data = aggregate_records(
                data_records,
                selected_dimension,
                selected_metric,
                selected_aggregation,
            )

            if category_chart_data:

                st.bar_chart(
                    category_chart_data,
                    x=selected_dimension,
                    y="Value",
                    x_label=selected_dimension,
                    y_label=(
                        f"{selected_aggregation} "
                        f"{selected_metric}"
                    ),
                    width="stretch",
                )

                st.dataframe(
                    category_chart_data,
                    width="stretch",
                    hide_index=True,
                )

            else:

                st.info(
                    "The selected fields do not "
                    "contain enough numeric data "
                    "for this chart."
                )


        date_columns = detect_date_columns(
            data_headers,
            data_records,
        )

        if (
            date_columns
            and numeric_columns
        ):

            st.markdown(
                "#### Trend Over Time"
            )

            trend_col1, trend_col2, trend_col3 = (
                st.columns(3)
            )

            with trend_col1:

                selected_date_column = st.selectbox(
                    "Date field",
                    options=date_columns,
                    key="dashboard_date",
                )

            with trend_col2:

                selected_trend_metric = st.selectbox(
                    "Trend metric",
                    options=numeric_columns,
                    key="dashboard_trend_metric",
                )

            with trend_col3:

                selected_trend_aggregation = st.selectbox(
                    "Trend aggregation",
                    options=[
                        "Average",
                        "Sum",
                        "Minimum",
                        "Maximum",
                    ],
                    key="dashboard_trend_aggregation",
                )

            trend_chart_data = aggregate_time_series(
                data_records,
                selected_date_column,
                selected_trend_metric,
                selected_trend_aggregation,
            )

            if trend_chart_data:

                st.line_chart(
                    trend_chart_data,
                    x=selected_date_column,
                    y="Value",
                    x_label=selected_date_column,
                    y_label=(
                        f"{selected_trend_aggregation} "
                        f"{selected_trend_metric}"
                    ),
                    width="stretch",
                )

            else:

                st.info(
                    "No usable time-series data "
                    "was found for the selected "
                    "date and metric."
                )


        if dimension_columns:

            st.markdown(
                "#### Category Distribution"
            )

            distribution_dimension = st.selectbox(
                "Category field",
                options=dimension_columns,
                key="dashboard_distribution_dimension",
            )

            distribution_data = aggregate_records(
                data_records,
                distribution_dimension,
                aggregation="Count",
            )

            if distribution_data:

                st.bar_chart(
                    distribution_data,
                    x=distribution_dimension,
                    y="Value",
                    x_label=distribution_dimension,
                    y_label="Record Count",
                    width="stretch",
                )

        default_data_company = st.session_state.get(
            "generated_company",
            "",
        )

        default_data_project = st.session_state.get(
            "generated_project_name",
            "",
        )

        analysis_company = st.text_input(
            "Company for Data Analysis",
            value=default_data_company,
            key="analysis_company",
        )

        analysis_project = st.text_input(
            "Project / Analysis Name",
            value=(
                default_data_project
                or "Business Data Analysis"
            ),
            key="analysis_project",
        )

        analysis_goal = st.text_area(
            "Analysis Goal (optional)",
            placeholder=(
                "Example: Identify operational bottlenecks, "
                "service trends, high-risk areas, and "
                "opportunities for process improvement."
            ),
            height=100,
            key="analysis_goal",
        )

        if st.button(
            "🧠 Analyze Business Data",
            use_container_width=True,
        ):

            if not data_records:
                st.warning(
                    "The uploaded file has no data rows to analyze."
                )

            elif not api_key:
                st.error(
                    "The OpenAI API key was not found."
                )

            else:
                profile_context = compact_profile_for_ai(
                    data_profile
                )

                sample_context = compact_records_for_ai(
                    data_headers,
                    data_records,
                )

                data_analysis_prompt = f"""
You are a senior Business Systems Analyst
with experience in operational analytics,
process improvement, KPI analysis, and
executive communication.

Analyze the uploaded business dataset using
ONLY the data profile and representative rows
provided below.

Company:
{analysis_company or "Not provided"}

Project / Analysis:
{analysis_project or "Business Data Analysis"}

User's Analysis Goal:
{analysis_goal or "Provide a general operational and business analysis."}

DATASET METRICS:

- Total Rows Read: {data_summary["rows"]}
- Total Columns: {data_summary["columns"]}
- Numeric Columns: {data_summary["numeric_columns"]}
- Missing Cells: {data_summary["missing_cells"]}
- Data Completeness: {data_summary["completeness"]}%

COLUMN PROFILE:

{profile_context}

REPRESENTATIVE DATA SAMPLE:

{sample_context}

Important limitations:

- The representative sample may not contain
  every row in the uploaded dataset.
- Do not claim a pattern exists across the
  entire dataset unless it is supported by
  the supplied profile or sample.
- Clearly label observations that require
  full-dataset validation.
- Do not invent company-specific facts.
- Do not infer protected or sensitive personal
  characteristics from the data.

Create the following sections:

## Executive Summary

Summarize the most decision-relevant findings.

## Data Quality Assessment

Discuss completeness, missing values,
field consistency, and analysis limitations.

## KPI and Operational Findings

Identify measurable observations supported
by the supplied data.

## Trends and Patterns

Identify notable trends, concentrations,
relationships, or unusual observations.
Distinguish confirmed findings from items
that require further validation.

## Business Risks

Identify operational, process, service,
financial, compliance, or data-quality risks
only when supported by the data.

## Process Improvement Opportunities

Recommend practical improvements connected
to the observed data.

## Recommended Actions

Provide 3 to 7 prioritized actions.
Number them sequentially.

## Suggested KPIs

Recommend useful KPIs that could be tracked
from this dataset or from closely related
business data.

## Follow-Up Analysis Questions

List the most valuable questions or additional
data needed for deeper analysis.

Use concise enterprise business-analysis
language suitable for stakeholders.
"""

                try:
                    data_analysis = generate_ai_text(
                        data_analysis_prompt,
                        "Analyzing business data and preparing recommendations...",
                    )

                    if not data_analysis:
                        st.error(
                            "OpenAI returned an empty response."
                        )
                    else:
                        st.session_state[
                            "generated_data_analysis"
                        ] = data_analysis

                        st.session_state[
                            "generated_data_filename"
                        ] = uploaded_business_file.name

                except Exception as error:
                    st.error(
                        "Business data analysis "
                        f"error: {error}"
                    )

        if st.session_state["generated_data_analysis"]:
            data_analysis = st.session_state[
                "generated_data_analysis"
            ]

            st.divider()

            st.success(
                "Business Data Analysis Generated"
            )

            st.markdown(data_analysis)

            st.subheader(
                "📥 Export Business Data Analysis"
            )

            show_download_buttons(
                "AI Business Data Analysis",
                data_analysis,
                analysis_company or "Not provided",
                analysis_project or "Business Data Analysis",
                "Business_Data_Analysis",
            )

    except (
        ValueError,
        zipfile.BadZipFile,
        ET.ParseError,
    ) as error:
        st.error(
            f"File reading error: {error}"
        )

    except Exception as error:
        st.error(
            "Unexpected file processing "
            f"error: {error}"
        )


# =========================================================
# COMPLETE PROJECT PACKAGE EXPORT
# =========================================================

st.divider()

st.header(
    "📦 Complete Project Package Export"
)

st.caption(
    "Combine the Business Analyst workflow into one "
    "professional project package for stakeholder review, "
    "portfolio demonstration, or project handoff."
)

package_requirements = st.session_state.get(
    "generated_requirements"
)

package_user_stories = st.session_state.get(
    "generated_user_stories"
)

package_test_cases = st.session_state.get(
    "generated_test_cases"
)

package_rtm = st.session_state.get(
    "generated_rtm"
)

package_executive_analysis = st.session_state.get(
    "generated_executive_analysis"
)

package_data_analysis = st.session_state.get(
    "generated_data_analysis"
)

package_data_filename = st.session_state.get(
    "generated_data_filename",
    "",
)

package_company = st.session_state.get(
    "generated_company",
    "",
)

package_project_name = st.session_state.get(
    "generated_project_name",
    "",
)

required_package_artifacts = {
    "BRD": bool(
        package_requirements
    ),
    "Jira Stories": bool(
        package_user_stories
    ),
    "QA Tests": bool(
        package_test_cases
    ),
    "RTM": bool(
        package_rtm
    ),
    "Executive Analysis": bool(
        package_executive_analysis
    ),
}

package_ready = all(
    required_package_artifacts.values()
)

status_columns = st.columns(
    len(
        required_package_artifacts
    )
)

for index, (
    artifact_label,
    artifact_ready,
) in enumerate(
    required_package_artifacts.items()
):

    with status_columns[index]:
        st.metric(
            artifact_label,
            (
                "Ready"
                if artifact_ready
                else "Missing"
            ),
        )

if package_data_analysis:
    st.success(
        "Optional Business Data Analysis is available "
        "and will be included in the package."
    )
else:
    st.info(
        "Business Data Analysis is optional. Generate one "
        "if you want it included in the final package."
    )

if not package_ready:
    missing_artifacts = [
        artifact_label
        for artifact_label, artifact_ready
        in required_package_artifacts.items()
        if not artifact_ready
    ]

    st.warning(
        "Complete these required artifacts before "
        "building the package: "
        + ", ".join(
            missing_artifacts
        )
        + "."
    )

if st.button(
    "📦 Build Complete Project Package",
    use_container_width=True,
    disabled=not package_ready,
):
    project_package_text = build_project_package_text(
        package_company,
        package_project_name,
        package_requirements,
        package_user_stories,
        package_test_cases,
        package_rtm,
        package_executive_analysis,
        package_data_analysis,
        package_data_filename,
    )

    st.session_state[
        "generated_project_package"
    ] = project_package_text


if st.session_state.get(
    "generated_project_package"
):
    project_package_text = st.session_state[
        "generated_project_package"
    ]

    st.success(
        "Complete Project Package Built"
    )

    st.markdown(
        "### Package Preview"
    )

    st.markdown(
        project_package_text
    )

    st.subheader(
        "📥 Download Complete Project Package"
    )

    show_download_buttons(
        "Complete Business Analysis Project Package",
        project_package_text,
        (
            package_company
            or "Not provided"
        ),
        (
            package_project_name
            or "Business Analysis Project"
        ),
        "Complete_Project_Package",
    )

    package_artifacts = [
        (
            "Business_Requirements_Document",
            package_requirements,
        ),
        (
            "Jira_User_Stories",
            package_user_stories,
        ),
        (
            "QA_Test_Cases",
            package_test_cases,
        ),
        (
            "Requirements_Traceability_Matrix",
            package_rtm,
        ),
        (
            "Executive_Project_Analysis",
            package_executive_analysis,
        ),
        (
            "Business_Data_Analysis",
            package_data_analysis,
        ),
    ]

    package_zip = create_project_package_zip(
        project_package_text,
        (
            package_company
            or "Not provided"
        ),
        (
            package_project_name
            or "Business Analysis Project"
        ),
        package_artifacts,
    )

    package_base_name = safe_filename(
        (
            package_project_name
            or "Business_Analysis_Project"
        )
    )

    st.download_button(
        label="🗂️ Download Complete ZIP Bundle",
        data=package_zip,
        file_name=(
            f"{package_base_name}"
            "_Complete_Project_Package.zip"
        ),
        mime="application/zip",
        use_container_width=True,
    )
